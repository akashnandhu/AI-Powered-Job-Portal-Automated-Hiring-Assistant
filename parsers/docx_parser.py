import docx
import os

def parse_docx(file_path):
    """
    Extracts text from a DOCX file using python-docx.
    Iterates through paragraphs and tables.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    text_content = []
    
    try:
        doc = docx.Document(file_path)
        
        # 1. Iterate through elements to maintain order (paragraphs and tables)
        for element in doc.element.body:
            # Check for Paragraph
            if element.tag.endswith('p'):
                para = [p for p in doc.paragraphs if p._element == element]
                if para:
                    text_content.append(para[0].text)
            
            # Check for Table
            elif element.tag.endswith('tbl'):
                table = [t for t in doc.tables if t._element == element]
                if table:
                    for row in table[0].rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        text_content.append(row_text)
                    
    except Exception as e:
        print(f"Error parsing DOCX {file_path}: {e}")
        return None

    return "\n".join(text_content)
